from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def generate_resume_docx(resume_data):
    document = Document()

    # ==========================
    # Header
    # ==========================

    name_paragraph = document.add_paragraph()
    name_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    name_run = name_paragraph.add_run(
        resume_data["full_name"]
    )

    name_run.bold = True
    name_run.font.size = Pt(20)

    contact_paragraph = document.add_paragraph()
    contact_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    contact_run = contact_paragraph.add_run(
        f'{resume_data["email"]} | {resume_data["phone"]}'
    )

    contact_run.font.size = Pt(10)

    # ==========================
    # Professional Summary
    # ==========================
    if resume_data.get("summary"):
        document.add_heading(
            "Professional Summary",
            level=2,
        )

        document.add_paragraph(
            resume_data["summary"]
        )

    # ==========================
    # Education
    # ==========================
    education_entries = resume_data.get(
        "education_entries",
        [],
)

    if education_entries:
        document.add_heading(
            "Education",
            level=2,
        )

        for entry in education_entries:

            institution = entry.get(
                "institution",
                "",
            )

            location = entry.get(
                "location",
                "",
            )

            qualification = entry.get(
                "qualification",
                "",
            )

            course = entry.get(
                "course",
                "",
            )

            start_year = entry.get(
                "start_year",
                "",
            )

            end_year = entry.get(
                "end_year",
                "",
            )

            grade = entry.get(
                "grade",
                "",
            )

            # Institution + Location
            institution_line = institution

            if institution and location:
                institution_line += (
                    f", {location}"
                )

            elif location:
                institution_line = location

            if institution_line:
                paragraph = document.add_paragraph()

                run = paragraph.add_run(
                    institution_line
                )

                run.bold = True

            # Education details
            details = ""

            # Years
            if start_year:
                details += start_year

            if end_year:
                if details:
                    details += f" – {end_year}"
                else:
                    details += end_year

            # Qualification + Course
            if qualification or course:

                if details:
                    details += " | "

                if qualification:
                    details += qualification

                if qualification and course:
                    details += " – "

                if course:
                    details += course

            # Grade
            if grade:

                if details:
                    details += " | "

                details += grade

            if details:
                document.add_paragraph(
                    details
                )

    # ==========================
    # Skills
    # ==========================
    if resume_data.get("skills"):
        document.add_heading(
            "Skills",
            level=2,
        )

        document.add_paragraph(
            resume_data["skills"]
        )

    # ==========================
    # Experience
    # ==========================

    experiences = resume_data.get(
        "experiences",
        [],
    )

    if experiences:
        document.add_heading(
            "Experience",
            level=2,
        )

        for experience_entry in experiences:

            company = experience_entry.get(
                "company",
                "",
            )

            job_title = experience_entry.get(
                "job_title",
                "",
            )

            duration = experience_entry.get(
                "duration",
                "",
            )

            description = experience_entry.get(
                "description",
                "",
            )

            # Company
            if company:
                company_paragraph = document.add_paragraph()

                company_run = company_paragraph.add_run(
                    company
                )

                company_run.bold = True

            # Job Title
            if job_title:
                document.add_paragraph(
                    job_title
                )

            # Duration
            if duration:
                document.add_paragraph(
                    duration
                )

            # Experience bullet points
            if description:

                for line in description.splitlines():

                    line = line.strip()

                    if not line:
                        continue

                    # Remove existing bullet characters
                    if line.startswith("•"):
                        line = line[1:].strip()

                    elif line.startswith("-"):
                        line = line[1:].strip()

                    # Remove Markdown formatting
                    line = line.replace("**", "")
                    line = line.replace("__", "")

                    document.add_paragraph(
                        line,
                        style="List Bullet",
                    )

    # ==========================
    # Projects
    # ==========================
    if resume_data.get("project_title"):
        document.add_heading(
            "Projects",
            level=2,
        )

        document.add_paragraph(
            resume_data["project_title"]
        )

        # Project technologies
        if resume_data.get("project_technologies"):
            document.add_paragraph(
                f'Technologies: '
                f'{resume_data["project_technologies"]}'
            )

        # Project description
        if resume_data.get("project_description"):
            document.add_paragraph(
                resume_data["project_description"]
            )

        # GitHub
        if resume_data.get("project_github"):
            document.add_paragraph(
                f'GitHub: '
                f'{resume_data["project_github"]}'
            )

    # ==========================
    # Certifications
    # ==========================
    if resume_data.get("certification_name"):
        document.add_heading(
            "Certifications",
            level=2,
        )

        certification = resume_data["certification_name"]

        if resume_data.get("certification_organization"):
            certification += (
                f' - '
                f'{resume_data["certification_organization"]}'
            )

        document.add_paragraph(
            certification
        )

        if resume_data.get("certification_issue_date"):
            document.add_paragraph(
                f'Issue Date: '
                f'{resume_data["certification_issue_date"]}'
            )

        if resume_data.get("certification_credential_id"):
            document.add_paragraph(
                f'Credential ID: '
                f'{resume_data["certification_credential_id"]}'
            )

    # ==========================
    # Languages
    # ==========================
    if resume_data.get("languages"):
        document.add_heading(
            "Languages",
            level=2,
        )

        document.add_paragraph(
            resume_data["languages"]
        )

    # ==========================
    # Save
    # ==========================
    output_path = "ResumeAI/generated_resume.docx"

    document.save(output_path)

    return output_path